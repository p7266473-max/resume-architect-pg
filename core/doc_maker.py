import io
import logging
from typing import Any

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.prompts import THEME_COLORS

logger = logging.getLogger("resume_architect")

def _clean(text: Any) -> str:
    """Removes XML-incompatible control characters."""
    if text is None:
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', str(text))

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz  # type: ignore
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            return "\n".join(page.get_text() for page in pdf)
    except ImportError:
        pass
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract  # type: ignore
        return pdfminer_extract(io.BytesIO(file_bytes))
    except ImportError:
        pass
    return ""

def _rgb(triple: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*triple)

def _add_hyperlink(paragraph: Any, url: str, text: str, color: RGBColor) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2]))
    rPr.append(clr)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def _separator(doc: Document, color: RGBColor, thickness_pt: float = 1.0) -> None:
    """Adds a beautiful native border line to divide resume sections cleanly."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(thickness_pt * 8)))  # sz is in 1/8 pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), "{:02X}{:02X}{:02X}".format(*color))
    pBdr.append(bottom)
    pPr.append(pBdr)

def _section_heading(doc: Document, text: str, color: RGBColor) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    
    r = h.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Arial'
    r.font.color.rgb = color
    
    # Add border
    pPr = h._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # 1pt border
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), "{:02X}{:02X}{:02X}".format(*color))
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_docx_bytes(data: dict, theme: str) -> bytes:
    """Build an ultra-premium, pixel-perfect DOCX resume."""
    colors = THEME_COLORS.get(theme, THEME_COLORS["Modern-Tech"])
    heading_color = _rgb(colors["heading"])
    sub_color = _rgb(colors["subheading"])
    accent_color = _rgb(colors["accent"])

    doc = Document()
    
    # Global Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia' # Premium serif for body, clean sans-serif for headers
    font.size = Pt(10)

    # Page Margins: 0.8 inch left/right means usable width is 8.5 - 1.6 = 6.9 inches
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    # ── HEADER: NAME & CONTACT ──────────────────────────────
    name = _clean(data.get("Name", "").strip() or "Your Name")
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    nr = name_para.add_run(name.upper())
    nr.bold = True
    nr.font.size = Pt(22)
    nr.font.name = 'Arial'
    nr.font.color.rgb = heading_color

    contact_parts = []
    if data.get("Location"): contact_parts.append(_clean(data.get("Location").strip()))
    if data.get("Phone"): contact_parts.append(_clean(data.get("Phone").strip()))
    if data.get("Email"): contact_parts.append(_clean(data.get("Email").strip()))
    if data.get("LinkedIn"): contact_parts.append(_clean(data.get("LinkedIn").strip()))

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(12)
        
        for i, part in enumerate(contact_parts):
            if "@" in part or "http" in part or "linkedin" in part.lower():
                _add_hyperlink(contact_para, part if "http" in part else f"mailto:{part}", part, accent_color)
            else:
                cr = contact_para.add_run(part)
                cr.font.size = Pt(9.5)
                cr.font.name = 'Arial'
                cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            
            if i < len(contact_parts) - 1:
                sep = contact_para.add_run("   |   ")
                sep.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── PROFESSIONAL SUMMARY ───────────────────────────────
    if data.get("Summary"):
        _section_heading(doc, "Professional Summary", heading_color)
        sp = doc.add_paragraph(_clean(data.get("Summary", "")))
        sp.paragraph_format.space_before = Pt(6)
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.line_spacing = 1.2
        sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in sp.runs:
            r.font.size = Pt(10)
            r.font.name = 'Georgia'

    # ── ACADEMIC & PERSONAL PROJECTS ────────────────────────────
    _section_heading(doc, "Academic & Personal Projects", heading_color)

    for exp in data.get("Experience", []):
        role_para = doc.add_paragraph()
        role_para.paragraph_format.space_before = Pt(8)
        role_para.paragraph_format.space_after = Pt(2)
        role_para.paragraph_format.keep_with_next = True
        
        # Set up a right-aligned tab stop at 6.9 inches for the date
        tab_stops = role_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.9), WD_TAB_ALIGNMENT.RIGHT)

        role = _clean(exp.get("Role", "Professional"))
        company = _clean(exp.get("Company", ""))
        duration = _clean(exp.get("Duration", ""))

        # Role & Company (Left)
        rr = role_para.add_run(role)
        rr.bold = True
        rr.font.size = Pt(11)
        rr.font.name = 'Arial'
        rr.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
        if company:
            cr = role_para.add_run(f"  |  {company}")
            cr.bold = True
            cr.font.size = Pt(11)
            cr.font.name = 'Arial'
            cr.font.color.rgb = sub_color

        # Date (Right aligned via tab)
        if duration:
            dr = role_para.add_run(f"\t{duration}")
            dr.bold = True
            dr.font.size = Pt(10)
            dr.font.name = 'Arial'
            dr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Achievements
        for ach in exp.get("Achievements", []):
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(3)
            bp.paragraph_format.space_after = Pt(3)
            bp.paragraph_format.line_spacing = 1.15
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.first_line_indent = Inches(-0.15)
            
            # Custom bullet point symbol
            bullet = bp.add_run("•  ")
            bullet.font.name = 'Arial'
            bullet.font.size = Pt(10)
            
            ar = bp.add_run(_clean(str(ach).strip()))
            ar.font.size = Pt(10)
            ar.font.name = 'Georgia'

    # ── KEY SKILLS ─────────────────────────────────────────
    skills = data.get("Skills", [])
    if skills:
        _section_heading(doc, "Core Competencies & Skills", heading_color)
        
        # Instead of a table, we use a beautifully formatted comma-separated block 
        # or grouped text to keep it 100% ATS proof and elegant.
        sp2 = doc.add_paragraph()
        sp2.paragraph_format.space_before = Pt(6)
        sp2.paragraph_format.space_after = Pt(10)
        sp2.paragraph_format.line_spacing = 1.3
        sp2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        skills_run = sp2.add_run(" • ".join([_clean(s) for s in skills]))
        skills_run.font.size = Pt(10)
        skills_run.font.name = 'Georgia'
        skills_run.font.bold = True
        skills_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── EDUCATION & CERTIFICATIONS ─────────────────────────
    edu_list = data.get("Education", [])
    if edu_list:
        _section_heading(doc, "Education & Certifications", heading_color)
        for edu in edu_list:
            ep = doc.add_paragraph()
            ep.paragraph_format.space_before = Pt(4)
            ep.paragraph_format.space_after = Pt(4)
            ep.paragraph_format.left_indent = Inches(0.15)
            
            # Format: split by " - " or "|" if generated that way, or just bold the whole line
            er = ep.add_run(_clean(str(edu)))
            er.bold = True
            er.font.size = Pt(10.5)
            er.font.name = 'Georgia'

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def generate_markdown(data: dict) -> str:
    """Generate a clean Markdown version of the resume."""
    lines: list[str] = []
    name = data.get("Name", "").strip()
    if name:
        lines.append(f"# {name}")
    contact = "  |  ".join(filter(None, [
        data.get("Email", ""), data.get("Phone", ""),
        data.get("LinkedIn", ""), data.get("Location", ""),
    ]))
    if contact:
        lines.append(contact)
    lines += ["", "---", "", "## Professional Summary", "", data.get("Summary", ""), "", "---", ""]

    if data.get("Experience"):
        lines.append("## Academic & Personal Projects")
        for exp in data["Experience"]:
            lines.append(f"\n### {exp.get('Role')}  ·  {exp.get('Company')}")
            lines.append(f"*{exp.get('Duration', '')}*")
            for ach in exp.get("Achievements", []):
                lines.append(f"- {ach}")

    lines += ["", "---", "", "## Key Skills", ""]
    if data.get("Skills"):
        lines.append(" • ".join(data["Skills"]))

    lines += ["", "---", "", "## Education & Certifications", ""]
    for edu in data.get("Education", []):
        lines.append(f"- **{edu}**")

    return "\n".join(lines)

def generate_ats_text(data: dict) -> str:
    """Generate a plain-text ATS-optimised version of the resume."""
    lines: list[str] = []
    name = data.get("Name", "").strip()
    if name:
        lines += [name.upper(), "=" * len(name)]
    for f in ("Email", "Phone", "LinkedIn", "Location"):
        v = data.get(f, "").strip()
        if v:
            lines.append(f"{f}: {v}")

    lines += ["", "PROFESSIONAL SUMMARY", "-" * 40, data.get("Summary", ""), ""]
    lines += ["ACADEMIC & PERSONAL PROJECTS", "-" * 40]
    for exp in data.get("Experience", []):
        lines.append(f"{exp.get('Role')} | {exp.get('Company')} | {exp.get('Duration', '')}")
        for ach in exp.get("Achievements", []):
            lines.append(f"  * {ach}")
        lines.append("")

    lines += ["KEY SKILLS", "-" * 40]
    lines.append(", ".join(data.get("Skills", [])))
    lines += ["", "EDUCATION & CERTIFICATIONS", "-" * 40]
    for edu in data.get("Education", []):
        lines.append(edu)

    return "\n".join(lines)
